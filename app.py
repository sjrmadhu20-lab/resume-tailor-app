import io
import json
import os
import re
import time
import zipfile
import docx
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor
from google import genai
from google.genai import types
import streamlit as st

st.set_page_config(
    page_title="Executive ATS Application Engine", page_icon="🎯", layout="wide"
)

# ==============================================================================
# 1. API CONFIGURATION
# ==============================================================================
api_key = st.secrets.get("GEMINI_API_KEY", os.environ.get("GEMINI_API_KEY", ""))

# ==============================================================================
# 2. MASTER KNOWLEDGE ARCHIVE (LOCKED BASELINE + DYNAMIC STORE)
# ==============================================================================
KNOWLEDGE_FILE = "custom_knowledge.json"

BASE_CAPABILITIES = {
    "commercial": (
        "Commercial & GTM Leadership ($100M+ P&L): Owned $100M+ annual FMCG"
        " revenue across GCC & India, directing 250+ distributors and 600+"
        " field sales teams across GT, MT, Wholesale, B2B, and Institutional"
        " channels. Spearheaded RTM redesign, distributor governance, trade"
        " margin economics, pricing/promotions, and Order-to-Cash optimization."
    ),
    "digital": (
        "Digital B2B2C Commerce & Omnichannel RTM: Founded and scaled Conektr"
        " (UAE's first digital FMCG distributor) to 8,000+ B2B retailers,"
        " managing 100+ brands and 2,000+ SKUs across Foods, Beverages, and"
        " Non-Food categories. Expanded into direct B2C commerce by launching"
        " the consumer app and proprietary BOSS loyalty engine (Buying,"
        " Operating, Selling & Saving), turning network grocers into"
        " fulfillment micro-hubs/dark stores. Built omnichannel ordering (App,"
        " Web, Conversational AI) with fintech-enabled payment rails."
    ),
    "transformation": (
        "Enterprise Transformation & Commercial Optimization: Directed"
        " multi-country RTM modernizations, DMS/ERP integrations and SFA"
        " deployments (Over 5000+ Users) for global CPG leaders (P&G, Nestlé,"
        " Haleon/GSK, Coca-Cola). Deployed AI route/beat optimization, AI-driven"
        " demand forecasting, and automated ordering—delivering a ~40% drop in"
        " logistics/admin costs, >30% reduction in outlet coverage costs, ~30%"
        " frontline sales productivity uplift, ~150% expansion in numeric"
        " distribution growth."
    ),
    "capability": (
        "Sales Capability & Training Leadership: Established and led regional"
        " sales training operations managing a team of certified trainers to"
        " design and deliver end-to-end sales induction and leadership"
        " curricula up to the Regional Sales Manager (RSM) level. Top-performer"
        " in consultative selling frameworks (including SPIN Selling), driving"
        " frontline execution rigor, distributor capability building, and"
        " institutionalized sales performance standards. Managed Train the"
        " Trainer, Soft skills and automation training."
    ),
    "entrepreneurship": (
        "Entrepreneurial Venture Scaling & Governance: Raised $15M in funding"
        " from DIFC VC, veteran FMCG executives (ex-Mondelēz President, BAT"
        " CFO) validating commercial credibility, and executed a successful"
        " strategic M&A exit to Al Maya Group ($1B+ conglomerate). Awarded UAE"
        " Golden Visa and USA O-1A (Extraordinary Ability); featured in"
        " Bloomberg, Gulf News, and Magnitt."
    ),
}

MASTER_STATIC = {
    "name": "MADHUSUDHANAN JANAKARAJAN (MADHU)",
    "contact": {
        "location": "Dubai, UAE",
        "phone": "+971 50 654 7858",
        "email": "sjrmadhu20@gmail.com",
        "email_url": "mailto:sjrmadhu20@gmail.com",
        "linkedin": "https://www.linkedin.com/in/madhusj/",
        "portfolio": "https://linktr.ee/M_S_J",
        "visas": "UAE Golden Visa | USA O-1A (Extraordinary Ability)",
    },
    "honors": [
        (
            "UAE Golden Visa – Recognized for national-scale entrepreneurship"
            " and digital commerce impact."
        ),
        (
            "USA O-1A Visa – Extraordinary Ability in FMCG and Digital"
            " Commerce."
        ),
        (
            "$15M+ VC funding & exit – Raised $15M+ and successfully exited"
            " Conektr to Al Maya Group."
        ),
        (
            "Featured in Gulf News, Bloomberg, Khaleej Times, Yahoo Finance,"
            " Magnitt, among others - https://linktr.ee/M_S_J"
        ),
    ],
    "education": [
        {
            "degree": "MBA (2006)",
            "details": (
                "Adam smith University, USA. [Remote, Airtel Sponsored program"
                " for top employees]"
            ),
        },
        {
            "degree": "Bachelor of Engineering (2001)",
            "details": (
                "Government College of Engineering (GEC), Tier 1 DOTE College,"
                " India"
            ),
        },
    ],
    "languages": (
        "English | Hindi | Tamil | Kannada | Telugu |   effective engagement"
        " with Arabic-speaking stakeholders."
    ),
    "interests": "Chess Player | Table Tennis Enthusiast | Regular 10K Runner",
    "tech_stack": {
        "AI, Automation & Conversational Commerce": (
            "Agentic Voice Bots (Vapi, ElevenLabs) | Conversational Commerce"
            " (Wati, Twilio, Infobip) | Workflow Automation (Make.com) | CRM &"
            " Marketing Automation (Klaviyo)"
        ),
        "Enterprise & Sales Systems": (
            "SAP (Sales & Distribution) | Oracle eCRM | Microsoft Dynamics |"
            " SFA / DMS platforms | ERP–CRMs API - integrations"
        ),
        "Digital Commerce & Product Delivery": (
            "WooCommerce | Magento | Mobile Apps (iOS, Android, Flutter) | Full"
            " SDLC ownership (Figma → Development → Launch)"
        ),
        "Data, Analytics & Optimization": (
            "Power BI | Alteryx | Tableau | Power Apps | Python scripting |"
            " Sales & trade analytics | Demand forecasting | Route & beat"
            " optimization"
        ),
        "Fintech & Payments": (
            "Stripe | PayPal | CCAvenue | Triterras | Tabby | Spotii (credit,"
            " payments, and trade finance integrations)"
        ),
    },
    "why_hire_me_parts": [
        ("A rare profile combining Core FMCG Operator ", False),
        ("+", True),
        (" Digital FMCG Disruption pioneer ", False),
        ("+", True),
        (" Enterprise Transformations (P&G, Coca-cola, GSK) ", False),
        ("+", True),
        (" 10+ International Markets (GCC, India, Africa, Asia) ", False),
        ("+", True),
        (" Successful Entrepreneurial $15M M&A Exit ", False),
        ("+", True),
        (
            (
                " Recipient of Global recognition for FMCG Contribution: O1A"
                " from USA & Golden Visa from UAE - as an extraordinary ability"
                " leader."
            ),
            False,
        ),
    ],
}

MASTER_DEEP_EXPERIENCE = """
CANDIDATE DEEP REPOSITORY & VERIFIED ACHIEVEMENTS:
- Category & Brand Portfolios:
  * Beverages & Energy (>50% Conektr GMV): Red Bull (18-month exclusivity contract with monthly trade fees), Coca-Cola (margin-backed exclusive contracts), Power Horse, PepsiCo, still/sparkling waters, sodas, and functional health drinks (Ornamin C, Vitane C).
  * Packaged Foods, Bakery & Snacking: Britannia biscuits/dairy P&L (200+ SKUs across GCC & India), Mondelez, Kellogg's, Kraft Heinz, and Nestlé.
  * Personal Care & Beauty: Extensive multi-category aggregation across Unilever (Dove, Sunsilk, Tresemme, Pond's, Vaseline), P&G (Pantene, Head & Shoulders, Olay), L'Oréal, and Colgate-Palmolive. (Note: Candidate specializes in mass/masstige high-volume retail CPG, NOT luxury boutique fashion).
  * Regulated Categories: Commercial distribution of international tobacco brands across UAE general trade (Marlboro, Dunhill, Chesterfield, Parliament, L&M, Rothmans).
  * Spirits: Commercial execution standards and digital pilots aligned with Bacardi's "Picture of Success".
  * Consumer Healthcare: SFA/DMS rollouts for GSK Consumer Healthcare / Haleon across MEA.
  * Telecom & Tech Sales: Airtel & Reliance GSM/CDMA SIMs, Broadband, web café internet trade models, and Samsung 'Slim' introduction in Karnataka. ADT Tyco enterprise electronic/smoke security tenders.
- Distributor Governance & Economics:
  * Governed 6 GCC master distributors and 200+ distributors in India for Britannia.
  * Techno-distribution philosophy: Introducing monthly unique technological interventions per distributor to demonstrate tangible ROI and drive distributor growth rather than only asking for numbers.
  * Command over NSV, Gross Margin (GM), trade spend ROI, pricing architecture, drop-size thresholds, and Order-to-Cash cycles. 30% margin improvements and 50% drop in supply chain complexities.
- Entrepreneurship & Operations:
  * Scaled Conektr to 8,000+ B2B stores, AED 50M GMV at 18% gross margin. Raised $15M from DIFC VC and C-suite FMCG leaders (ex-President Mondelez, ex-CFO BAT) with 25% to 400% returns on M&A exit to Al Maya Group.
  * Overcame two-sided marketplace chicken-and-egg hurdle by partnering with master supermarket for full 10K SKU master catalog, building immediate retail density before transitioning directly to brand principals.
  * Resolved field sales resistance to digital self-ordering by aligning incentive plans, freeing reps to focus on high-commission consultative category development.
  * Lean build: Deployed Magento, Flutter, Dynamics 365, VPN telephony on <10% standard capital. Built offshore delivery hubs cutting operational costs by 66%.
  * Last-mile logistics: 2-wheeler urban fleet cutting fulfillment costs by ~50%.
- Omnichannel & Product Management:
  * Full SDLC ownership (BRD/PRD, Figma, QA, Launch) positioning platform value at $27M.
  * Proprietary BOSS engine (Buying, Operating, Selling & Saving) turning network grocers into micro-fulfillment dark stores.
  * White-labeled B2C ordering app with Click & Collect, Ship-from-Store, and in-store 'Scan & Go' self-checkout with automated audio alerts.
  * Full-funnel digital marketing (Meta/Google Ads, SEO, SEM, Klaviyo) reducing CAC by ~40%.
- Transformation, Organization Design & Field Rigor:
  * Personally deployed mobile SFA on ground with P&G distributor networks in Kenya.
  * Enterprise SFA/DMS/ERP modernizations (SAP SD, Dynamics 365, Oracle e-CRM) for 5,000+ sales users across Britannia and Ivy Mobility.
  * Computer vision / Image Recognition (IR) audits: Automated shelf-share analytics, stock-out detection, and predictive reordering linked to Tableau/Qlik.
  * Macro Organization Design (OD) & Organization Effectiveness (OE) partner counseling C-suite leaders. Capability Champion networks, Train-the-Trainer (TTT), BMI (Business Measurement Index) distributor audits, and Kirkpatrick 70:20:10 learning frameworks.
"""

def load_custom_knowledge():
    """Loads dynamically added Q&A knowledge from disk."""
    if os.path.exists(KNOWLEDGE_FILE):
        try:
            with open(KNOWLEDGE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []

def save_custom_knowledge(new_entries):
    """Appends validated Q&A context to the dynamic knowledge repository."""
    existing = load_custom_knowledge()
    existing.extend(new_entries)
    try:
        with open(KNOWLEDGE_FILE, "w", encoding="utf-8") as f:
            json.dump(existing, f, indent=2, ensure_ascii=False)
        return True
    except Exception:
        return False

def get_full_knowledge_context():
    """Combines static deep achievements with persistent Q&A inputs."""
    custom_items = load_custom_knowledge()
    custom_text = "\n".join([f"- USER RECORDED FACT: Q: {item.get('q', '')} | A: {item.get('a', '')}" for item in custom_items])
    return MASTER_DEEP_EXPERIENCE + ("\n\nDYNAMICALLY STORED CUSTOM FACTS:\n" + custom_text if custom_text else "")

def clean_ai_generated_text(text):
    """Deterministic safeguard to prevent hallucinated phrasing and keep executive metrics clean."""
    if not isinstance(text, str):
        return text
    replacements = [
        (r"\bthirty-six-degree\b", "360°"),
        (r"\bthirty six degree\b", "360°"),
        (r"\bthree hundred and sixty degree\b", "360°"),
        (r"\b360 degree\b", "360°"),
        (r"\b360-degree\b", "360°"),
        (r"\btwenty-three years\b", "23+ years"),
        (r"\bone hundred million dollars\b", "$100M+"),
        (r"\beight thousand\b", "8,000+"),
        (r"\bfifteen million\b", "$15M+"),
    ]
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text, flags=re.IGNORECASE)
    return text

def sanitize_json_payload(data):
    """Recursively cleans all strings in the AI output dictionary."""
    if isinstance(data, dict):
        return {k: sanitize_json_payload(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [sanitize_json_payload(item) for item in data]
    elif isinstance(data, str):
        return clean_ai_generated_text(data)
    return data

def add_hyperlink(paragraph, url, text, color_rgb="004B87", underline=True, font_size_pt=10, is_highlighted=False):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = parse_xml(f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>')
    new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    rPr = parse_xml(f'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    rPr.append(parse_xml(f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="Calibri" w:hAnsi="Calibri"/>'))
    val_sz = int(font_size_pt * 2)
    rPr.append(parse_xml(f'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{val_sz}"/>'))
    rPr.append(parse_xml(f'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{color_rgb}"/>'))
    if underline:
        rPr.append(parse_xml(f'<w:u xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="single"/>'))
    if is_highlighted:
        rPr.append(parse_xml(r'<w:highlight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="yellow"/>'))
        
    new_run.append(rPr)
    new_run.append(parse_xml(f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{text}</w:t>'))
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# ==============================================================================
# 3. WORD RESUME ENGINE (LOCKED V1 BUILD)
# ==============================================================================
def populate_resume_document(doc, tailored_data, highlight_changes=False):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def apply_xml_spacing(p, before_pt=0, after_pt=8, line_twips=278):
        pPr = p._p.get_or_add_pPr()
        spPr = parse_xml(f'<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:before="{int(before_pt*20)}" w:after="{int(after_pt*20)}" w:line="{line_twips}" w:lineRule="auto"/>')
        pPr.append(spPr)

    def add_heading(title, space_before=0, space_after=8, line_border_above=False, is_multiple=False, is_underline=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if is_multiple:
            apply_xml_spacing(p, before_pt=space_before, after_pt=space_after, line_twips=278)
        else:
            apply_xml_spacing(p, before_pt=space_before, after_pt=space_after, line_twips=240)
        
        if line_border_above:
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:top w:val="single" w:sz="6" w:space="4" w:color="000000"/>'
                             r'</w:pBdr>')
            pPr.append(pBdr)
            
        r = p.add_run(title.upper() if title != "LANGUAGES & INTERESTS :" else title)
        r.bold = True
        r.underline = is_underline
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # ---------------- PAGE 1 (LOCKED V1) ----------------
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_xml_spacing(p_name, before_pt=0, after_pt=0, line_twips=278)
    r_name = p_name.add_run(MASTER_STATIC['name'])
    r_name.bold = True
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(12)

    f1 = tailored_data.get("header_focus_1", "Sales & Distribution Transformation Director")
    f2 = tailored_data.get("header_focus_2", "Beauty & Personal Care Experience")
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_xml_spacing(p_sub, before_pt=0, after_pt=8, line_twips=278)
    
    r_f1 = p_sub.add_run(f1)
    r_f1.bold = True
    r_f1.font.name = 'Calibri'
    r_f1.font.size = Pt(9)
    if highlight_changes:
        r_f1.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
        
    r_mid = p_sub.add_run(" | FMCG | GTM & Omnichannel Leader | ")
    r_mid.bold = True
    r_mid.font.name = 'Calibri'
    r_mid.font.size = Pt(9)
    
    r_f2 = p_sub.add_run(f2)
    r_f2.bold = True
    r_f2.font.name = 'Calibri'
    r_f2.font.size = Pt(9)
    if highlight_changes:
        r_f2.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

    c = MASTER_STATIC['contact']
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_xml_spacing(p_contact, before_pt=0, after_pt=8, line_twips=278)
    
    r_c1 = p_contact.add_run(f"{c['location']} | {c['phone']} | ")
    r_c1.font.name = 'Calibri'
    r_c1.font.size = Pt(10)
    add_hyperlink(p_contact, c['email_url'], c['email'], color_rgb="004B87", underline=True, font_size_pt=10)
    
    r_br1 = p_contact.add_run("\n")
    r_br1.font.name = 'Calibri'
    r_br1.font.size = Pt(10)
    
    add_hyperlink(p_contact, c['linkedin'], c['linkedin'], color_rgb="004B87", underline=True, font_size_pt=10)
    r_c2_mid = p_contact.add_run(" | Portfolio: ")
    r_c2_mid.font.name = 'Calibri'
    r_c2_mid.font.size = Pt(10)
    add_hyperlink(p_contact, c['portfolio'], c['portfolio'], color_rgb="004B87", underline=True, font_size_pt=10)
    
    r_br2 = p_contact.add_run("\n")
    r_br2.font.name = 'Calibri'
    r_br2.font.size = Pt(10)
    
    r_c3_lbl = p_contact.add_run("Visa Status: ")
    r_c3_lbl.bold = True
    r_c3_lbl.font.name = 'Calibri'
    r_c3_lbl.font.size = Pt(10)
    r_c3_val = p_contact.add_run(c['visas'])
    r_c3_val.font.name = 'Calibri'
    r_c3_val.font.size = Pt(10)

    add_heading("EXECUTIVE SUMMARY", space_before=0, space_after=8, line_border_above=False, is_multiple=True)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    apply_xml_spacing(sp, before_pt=0, after_pt=6, line_twips=278)
    r_sum = sp.add_run(tailored_data.get("executive_summary", ""))
    r_sum.font.name = 'Calibri'
    r_sum.font.size = Pt(10)
    if highlight_changes:
        r_sum.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

    add_heading("EXECUTIVE CAPABILITIES & IMPACT HIGHLIGHTS", space_before=2, space_after=8, line_border_above=True, is_multiple=False)
    for cap in tailored_data.get("capabilities", []):
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Inches(0.20)
        cp.paragraph_format.first_line_indent = Inches(-0.25)
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_xml_spacing(cp, before_pt=0, after_pt=4.5, line_twips=240)
        
        r_bullet = cp.add_run("•\t")
        r_bullet.font.name = 'Calibri'
        r_bullet.font.size = Pt(10)
        
        parts = cap.split(":", 1)
        if len(parts) == 2:
            r_bold = cp.add_run(parts[0] + ":")
            r_bold.bold = True
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(10)
            r_body = cp.add_run(parts[1])
            r_body.font.name = 'Calibri'
            r_body.font.size = Pt(10)
        else:
            r_body = cp.add_run(cap)
            r_body.font.name = 'Calibri'
            r_body.font.size = Pt(10)

    add_heading("HONORS & RECOGNITION", space_before=2, space_after=8, line_border_above=True, is_multiple=False)
    for idx, h in enumerate(MASTER_STATIC['honors']):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.45)
        p.paragraph_format.first_line_indent = Inches(-0.20)
        is_last = (idx == len(MASTER_STATIC['honors']) - 1)
        apply_xml_spacing(p, before_pt=0, after_pt=6 if is_last else 0, line_twips=240)
        
        r_bullet = p.add_run("•\t")
        r_bullet.font.name = 'Calibri'
        r_bullet.font.size = Pt(10)
        
        if "https://" in h:
            parts = h.split(" - ")
            r_prefix = p.add_run(parts[0] + " - ")
            r_prefix.font.name = 'Calibri'
            r_prefix.font.size = Pt(10)
            add_hyperlink(p, parts[1].strip(), parts[1].strip(), color_rgb="004B87", underline=True, font_size_pt=10)
        else:
            r_t = p.add_run(h)
            r_t.font.name = 'Calibri'
            r_t.font.size = Pt(10)

    add_heading("EDUCATION", space_before=0, space_after=8, line_border_above=False, is_multiple=False)
    for idx, edu in enumerate(MASTER_STATIC['education']):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.45)
        p.paragraph_format.first_line_indent = Inches(-0.20)
        is_last = (idx == len(MASTER_STATIC['education']) - 1)
        apply_xml_spacing(p, before_pt=0, after_pt=6 if is_last else 0, line_twips=240)
        
        r_bullet = p.add_run("•\t")
        r_bullet.font.name = 'Calibri'
        r_bullet.font.size = Pt(10)
        
        r_bp = p.add_run(edu['degree'] + " – ")
        r_bp.bold = True
        r_bp.font.name = 'Calibri'
        r_bp.font.size = Pt(10)
        r_t = p.add_run(edu['details'])
        r_t.font.name = 'Calibri'
        r_t.font.size = Pt(10)

    add_heading("LANGUAGES & INTERESTS :", space_before=0, space_after=8, line_border_above=False, is_multiple=False)
    p_lang1 = doc.add_paragraph()
    p_lang1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_lang1.paragraph_format.left_indent = Inches(0.45)
    p_lang1.paragraph_format.first_line_indent = Inches(-0.20)
    apply_xml_spacing(p_lang1, before_pt=0, after_pt=0, line_twips=240)
    r_bullet_l1 = p_lang1.add_run("•\t")
    r_bullet_l1.font.name = 'Calibri'
    r_bullet_l1.font.size = Pt(10)
    r_l1 = p_lang1.add_run(MASTER_STATIC['languages'])
    r_l1.font.name = 'Calibri'
    r_l1.font.size = Pt(10)

    p_lang2 = doc.add_paragraph()
    p_lang2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_lang2.paragraph_format.left_indent = Inches(0.45)
    p_lang2.paragraph_format.first_line_indent = Inches(-0.20)
    apply_xml_spacing(p_lang2, before_pt=0, after_pt=0, line_twips=240)
    r_bullet_l2 = p_lang2.add_run("•\t")
    r_bullet_l2.font.name = 'Calibri'
    r_bullet_l2.font.size = Pt(10)
    r_l2 = p_lang2.add_run(MASTER_STATIC['interests'])
    r_l2.font.name = 'Calibri'
    r_l2.font.size = Pt(10)

    # ---------------- PAGE 2 (LOCKED V1) ----------------
    doc.add_page_break()

    add_heading("PROFESSIONAL EXPERIENCE", space_before=0, space_after=8, line_border_above=False, is_multiple=False)
    
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    
    tblPr = table._tbl.tblPr
    tblpPr = parse_xml(r'<w:tblpPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:tblpX="-274" w:tblpY="0"/>')
    tblPr.append(tblpPr)
    
    col_widths = [Inches(2.63), Inches(2.63), Inches(2.63)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Contextual dynamic headings for the 3 columns
    h0 = tailored_data.get("exp_col_header_1", "Traditional FMCG Operator")
    h1 = tailored_data.get("exp_col_header_2", "Digital FMCG Distribution")
    h2 = tailored_data.get("exp_col_header_3", "Distribution Transformation")
    hdr_titles = [h0, h1, h2]

    for i, title in enumerate(hdr_titles):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_xml_spacing(p, before_pt=3, after_pt=3, line_twips=240)
        r = p.add_run(title)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="DCE6F1"/>'))

    def populate_cell_content(cell, item_list):
        cell.text = ""
        for idx, item in enumerate(item_list):
            p = cell.add_paragraph()
            apply_xml_spacing(p, before_pt=item.get("space_before", 0), after_pt=item.get("space_after", 2), line_twips=220)
            
            if item.get("is_bullet", False):
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.first_line_indent = Inches(-0.15)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r_b = p.add_run("•\t")
                r_b.font.name = 'Calibri'
                r_b.font.size = Pt(10)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(item["text"])
            r.bold = item.get("bold", False)
            r.italic = item.get("italic", False)
            r.font.name = 'Calibri'
            r.font.size = Pt(item.get("size", 10))
            if highlight_changes and item.get("highlight", False):
                r.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

    c0_items = [
        {"text": "Britannia Industries Ltd | 2007 – 2011", "bold": True, "size": 10, "space_before": 2},
        {"text": "Regional Sales Head – GCC", "bold": True, "size": 10},
        {"text": "Regional Sales & Capability Head- India", "bold": True, "size": 10, "space_after": 4},
        {"text": "Owned $100M+ P&L across GCC (Saudi Arabia, UAE, Kuwait, Oman, Bahrain, Qatar) & South India.", "is_bullet": True, "size": 10},
        {"text": "Directed 250+ distributor networks & 600+ frontline sales staff across GT, MT, wholesale, and institutional trade.", "is_bullet": True, "size": 10},
        {"text": "Spearheaded Britannia's 1st national SFA rollout (1,000+ users), transforming legacy trade into performance-managed selling.", "is_bullet": True, "size": 10},
        {"text": "Delivered ~30% numeric distribution growth, increased LPC to ~120%, and cut sales admin costs by ~30%.", "is_bullet": True, "size": 10},
        {"text": "Turnaround RSM GCC: achieved record monthly sales for 3 consecutive months (Best Employee Award from Group MD).", "is_bullet": True, "size": 10, "space_after": 3},
        {"text": "Airtel | Reliance | Tyco | 2001 – 2007", "bold": True, "size": 10, "space_before": 5},
        {"text": "Commercial & Training Roles –", "bold": True, "size": 10, "space_after": 4},
        {"text": "Built foundations in frontline trade execution, journey planning, and merchandiser enablement in telecom & enterprise security.", "is_bullet": True, "size": 10},
        {"text": "Deployed capability training (SPIN selling) & integrated Oracle e-CRM & LMS infrastructure at scale.", "is_bullet": True, "size": 10}
    ]

    c0_extra = tailored_data.get("column_1_extra_bullet", "")
    if c0_extra and c0_extra.strip():
        c0_items.insert(7, {"text": c0_extra.strip(), "is_bullet": True, "size": 10, "highlight": True})

    conektr_cat = tailored_data.get("conektr_category_bullet", "Deep FMCG Category Aggregation: Scaled multi-category catalogs across ambient, packaged food, and consumer goods portfolios.")
    c1_items = [
        {"text": "Digital FMCG Principal / Distributor", "bold": True, "size": 10, "space_before": 2},
        {"text": "Conektr Tech Global Ltd | UAE & India", "bold": True, "size": 10, "space_after": 4},
        {"text": "Chief Executive Officer & Founder", "bold": True, "size": 10},
        {"text": "May 2016 – Aug 2024", "size": 10, "space_after": 4},
        {"text": "Founded UAE’s 1st Digital FMCG Principal-Distributor serving 8,000+ retailers (2,000+ MAU) & 100+ brands.", "is_bullet": True, "size": 10},
        {"text": conektr_cat, "is_bullet": True, "size": 10, "highlight": True},
        {"text": "Owned full P&L, trade terms, warehousing, last-mile delivery, trade credit, and collections.", "is_bullet": True, "size": 10},
        {"text": "Built app/web/WhatsApp self-ordering engine scaling annual GMV from zero to ~AED 50M (~$13.6M) at ~18% gross margin.", "is_bullet": True, "size": 10},
        {"text": "Cut coverage cost by >50% and improved field execution productivity by ~150% vs traditional trade.", "is_bullet": True, "size": 10},
        {"text": "Deployed Dynamics 365 + Power BI and AI route optimization, cutting logistics costs by ~40%.", "is_bullet": True, "size": 10},
        {"text": "Raised ~$15M from C-suite FMCG leaders; executed M&A exit to Al Maya Group ($1B+ retail conglomerate).", "is_bullet": True, "size": 10}
    ]

    c2_items = [
        {"text": "Post Exit –", "size": 10, "space_before": 2, "space_after": 3},
        {"text": "Transformation Advisor (Director)", "bold": True, "size": 10},
        {"text": "TransCPG Inc. &", "bold": True, "size": 10},
        {"text": "FieldAssist | 2025 – Present", "bold": True, "size": 10, "space_after": 4},
        {"text": "Board Member guiding global operations scaling & platform build across FMCG principals & distributors.", "is_bullet": True, "size": 10},
        {"text": "Advising CPG leaders on modernizing RTM & SAP/Oracle SFA/DMS integrations, driving ~150% coverage growth.", "is_bullet": True, "size": 10},
        {"text": "Built Bid2Bill AI/Voice-bot & WhatsApp B2B2C bidding platform, cutting CAC by ~40% with 4x engagement.", "is_bullet": True, "size": 10, "space_after": 3},
        {"text": "Business Head – MEA", "bold": True, "size": 10, "space_before": 5},
        {"text": "Ivy Mobility Pte Ltd | 2011 – 2016", "bold": True, "size": 10, "space_after": 4},
        {"text": "Built MEA setup from scratch into 2nd largest global setup ($10M+ pipeline across 10+ countries).", "is_bullet": True, "size": 10},
        {"text": "Won 22 enterprise logos: Haleon/GSK, P&G, Nestlé, Coca-Cola, Mars, Red Bull, BAT, and AKI Group.", "is_bullet": True, "size": 10},
        {"text": "Personally led on-ground field deployment of mobile SFA for P&G distributor networks in Kenya.", "is_bullet": True, "size": 10},
        {"text": "Deployed Cloud SaaS SFA/DMS to 3,000+ sales users, driving post-implementation adoption and trade ROI.", "is_bullet": True, "size": 10}
    ]

    c1_extra = tailored_data.get("column_2_extra_bullet", "")
    if c1_extra and c1_extra.strip():
        c1_items.insert(7, {"text": c1_extra.strip(), "is_bullet": True, "size": 10, "highlight": True})

    c2_extra = tailored_data.get("column_3_extra_bullet", "")
    if c2_extra and c2_extra.strip():
        c2_items.insert(4, {"text": c2_extra.strip(), "is_bullet": True, "size": 10, "highlight": True})

    populate_cell_content(table.rows[1].cells[0], c0_items)
    populate_cell_content(table.rows[1].cells[1], c1_items)
    populate_cell_content(table.rows[1].cells[2], c2_items)

    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tblBorders>'
    )
    table._tbl.tblPr.append(tblBorders)

    add_heading("TECHNOLOGY STACK & DIGITAL ARCHITECTURE:", space_before=8, space_after=8, line_border_above=False, is_multiple=False)
    for category, stack in MASTER_STATIC['tech_stack'].items():
        tp = doc.add_paragraph()
        tp.paragraph_format.left_indent = Inches(0.20)
        tp.paragraph_format.first_line_indent = Inches(-0.25)
        tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_xml_spacing(tp, before_pt=0, after_pt=4, line_twips=240)
        
        r_b = tp.add_run("•\t")
        r_b.font.name = 'Calibri'
        r_b.font.size = Pt(10)
        
        r_cat = tp.add_run(f"{category}: ")
        r_cat.bold = True
        r_cat.font.name = 'Calibri'
        r_cat.font.size = Pt(10)
        
        r_st = tp.add_run(stack)
        r_st.font.name = 'Calibri'
        r_st.font.size = Pt(10)

    add_heading("WHY HIRE ME", space_before=8, space_after=4, line_border_above=False, is_multiple=False, is_underline=True)
    
    p_why = doc.add_paragraph()
    p_why.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    apply_xml_spacing(p_why, before_pt=0, after_pt=6, line_twips=240)
    
    for text_segment, is_plus in MASTER_STATIC['why_hire_me_parts']:
        r_part = p_why.add_run(text_segment)
        r_part.font.name = 'Calibri'
        r_part.font.size = Pt(10)
        if is_plus:
            r_part.bold = True
            r_part.font.color.rgb = RGBColor(0x00, 0xB0, 0xF0)

def create_master_resume_docx(tailored_data, highlight_changes=False):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.40)
        section.bottom_margin = Inches(0.40)
        section.left_margin = Inches(0.50)
        section.right_margin = Inches(0.50)
    populate_resume_document(doc, tailored_data, highlight_changes)
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

# ==============================================================================
# 4. WORD COVER, MATRIX & COMBINED PACK BUILDER (.DOCX)
# ==============================================================================
def populate_cover_letter_docx_page(doc, cover_data):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(10)
    r_t = p_title.add_run("COVER LETTER")
    r_t.bold = True
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(14)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(6)
    p_sub.paragraph_format.space_after = Pt(8)
    r_sb = p_sub.add_run(f"Subject: {cover_data.get('subject_line', '')}")
    r_sb.bold = True
    r_sb.font.name = 'Calibri'
    r_sb.font.size = Pt(11)

    p_d = doc.add_paragraph("Dear Hiring Team,")
    p_d.paragraph_format.space_before = Pt(6)
    p_d.paragraph_format.space_after = Pt(8)

    p_p1 = doc.add_paragraph(cover_data.get("cover_para_1", ""))
    p_p1.paragraph_format.space_after = Pt(8)
    p_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_p2 = doc.add_paragraph(cover_data.get("cover_para_2", ""))
    p_p2.paragraph_format.space_after = Pt(8)
    p_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_kh = doc.add_paragraph()
    p_kh.paragraph_format.space_after = Pt(6)
    r_kh = p_kh.add_run("Key highlights of what I bring to this mandate include:")
    r_kh.bold = True
    r_kh.font.name = 'Calibri'
    r_kh.font.size = Pt(11)

    for b in cover_data.get("cover_bullets", []):
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.left_indent = Inches(0.5)
        bp.paragraph_format.right_indent = Inches(0.2)
        bp.paragraph_format.space_after = Pt(6)
        bp.paragraph_format.line_spacing = 1.15
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        parts = b.split(":", 1)
        if len(parts) == 2:
            r_head = bp.add_run(parts[0] + ": ")
            r_head.bold = True
            r_head.font.name = 'Calibri'
            r_head.font.size = Pt(10.5)
            r_tail = bp.add_run(parts[1].strip())
            r_tail.font.name = 'Calibri'
            r_tail.font.size = Pt(10.5)
        else:
            r_b = bp.add_run(b)
            r_b.font.name = 'Calibri'
            r_b.font.size = Pt(10.5)

    p_cl = doc.add_paragraph(cover_data.get("cover_para_closing", ""))
    p_cl.paragraph_format.space_before = Pt(6)
    p_cl.paragraph_format.space_after = Pt(8)
    p_cl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_sign = doc.add_paragraph()
    p_sign.paragraph_format.space_before = Pt(8)
    r_s0 = p_sign.add_run("Sincerely,\n")
    r_s0.font.name = 'Calibri'
    r_s1 = p_sign.add_run("Madhusudhanan Janakarajan (Madhu)\n")
    r_s1.bold = True
    r_s1.font.name = 'Calibri'
    r_s2 = p_sign.add_run("+971 50 654 7858 | sjrmadhu20@gmail.com")
    r_s2.font.name = 'Calibri'

def populate_match_matrix_docx_page(doc, cover_data):
    p_mtitle = doc.add_paragraph()
    p_mtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mtitle.paragraph_format.space_after = Pt(8)
    r_mt = p_mtitle.add_run("MATCH MATRIX")
    r_mt.bold = True
    r_mt.font.name = 'Calibri'
    r_mt.font.size = Pt(14)

    matrix_items = cover_data.get("matrix_items", [])
    table = doc.add_table(rows=len(matrix_items) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0].cells[0].width = Inches(2.5)
    table.rows[0].cells[1].width = Inches(5.0)

    cell_0 = table.rows[0].cells[0]
    cell_1 = table.rows[0].cells[1]
    
    p_h0 = cell_0.paragraphs[0]
    p_h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h0 = p_h0.add_run("Target Job Requirement / Focus Domain")
    r_h0.bold = True
    r_h0.font.name = 'Calibri'
    r_h0.font.size = Pt(10.5)

    p_h1 = cell_1.paragraphs[0]
    p_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h1 = p_h1.add_run("How I Match (Evidence & Track Record)")
    r_h1.bold = True
    r_h1.font.name = 'Calibri'
    r_h1.font.size = Pt(10.5)

    for idx, item in enumerate(matrix_items):
        row_cells = table.rows[idx + 1].cells
        row_cells[0].width = Inches(2.5)
        row_cells[1].width = Inches(5.0)
        
        p0 = row_cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after = Pt(3)
        r_rt = p0.add_run(item.get('requirement_title', ''))
        r_rt.bold = True
        r_rt.font.name = 'Calibri'
        r_rt.font.size = Pt(9.5)
        
        p1 = row_cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        r_mt = p1.add_run(item.get('match_desc', ''))
        r_mt.font.name = 'Calibri'
        r_mt.font.size = Pt(9.5)

    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        r'</w:tblBorders>'
    )
    table._tbl.tblPr.append(tblBorders)

def create_combined_application_docx(cover_data, tailored_data):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.40)
        section.bottom_margin = Inches(0.40)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    populate_cover_letter_docx_page(doc, cover_data)
    doc.add_page_break()
    populate_resume_document(doc, tailored_data, highlight_changes=False)
    doc.add_page_break()
    populate_match_matrix_docx_page(doc, cover_data)
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def create_master_application_zip(comb_docx, review_docx, clean_docx):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr('1_Complete_Application_Set_Cover_Resume_Matrix.docx', comb_docx)
        zip_file.writestr('2_Madhusudhanan_Janakarajan_Resume_Highlighted_Review.docx', review_docx)
        zip_file.writestr('3_Madhusudhanan_Janakarajan_Resume_Clean.docx', clean_docx)
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def rebuild_all_documents():
    """Helper to recompile all 3 docx files and zip bundle when state changes."""
    tailored_data = st.session_state["tailored_data"]
    cover_data = st.session_state["cover_data"]
    clean_docx = create_master_resume_docx(tailored_data, highlight_changes=False)
    review_docx = create_master_resume_docx(tailored_data, highlight_changes=True)
    comb_docx = create_combined_application_docx(cover_data, tailored_data)
    master_zip = create_master_application_zip(comb_docx, review_docx, clean_docx)

    st.session_state["comb_docx"] = comb_docx
    st.session_state["review_docx"] = review_docx
    st.session_state["clean_docx"] = clean_docx
    st.session_state["master_zip"] = master_zip

# ==============================================================================
# 5. STREAMLIT FRONTEND & ENGINE CONTROLLER
# ==============================================================================
st.title("🎯 Executive ATS Resume & Application Engine")
st.caption("Contextual Track Routing • Intelligent Gap Questioning • In-Place Revisions • Word (.docx) Suite")

with st.sidebar:
    st.header("⚡ System Status")
    if api_key:
        st.success("🟢 Gemini AI Engine: Active (gemini-3.6-flash)")
    else:
        st.error("🔴 AI Engine Key Missing (Set GEMINI_API_KEY in Secrets)")
    
    st.markdown("---")
    stored_facts = load_custom_knowledge()
    st.write(f"📚 **Dynamic Custom Facts Stored:** {len(stored_facts)}")
    if st.button("🔄 Reset / Clear Cached Session"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

col1, col2 = st.columns([1.1, 0.9])

with col1:
    st.subheader("1. Job Inputs & Context")
    job_desc = st.text_area(
        "Target Job Description (JD):",
        height=220,
        placeholder="Paste target Job Description here...",
    )

    st.markdown("##### Special Instructions & Context (Optional)")
    st.caption("Casual notes or target company nuances are automatically formalized and contextualized.")

    st.components.v1.html(
        """
        <div style="font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; margin-bottom: 6px;">
            <button id="micBtn" onclick="toggleDictation()" style="
                background-color: #2563EB;
                color: white;
                border: none;
                padding: 7px 14px;
                font-size: 13px;
                font-weight: 600;
                border-radius: 6px;
                cursor: pointer;
                display: inline-flex;
                align-items: center;
                gap: 6px;
            ">🎙️ Click to Speak Instructions</button>
            <span id="status" style="font-size: 12px; color: #4B5563; margin-left: 10px; font-weight: 500;"></span>
        </div>
        <script>
            var recognizing = false;
            var recognition;
            if ('webkitSpeechRecognition' in window || 'SpeechRecognition' in window) {
                var SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
                recognition = new SpeechRecognition();
                recognition.continuous = true;
                recognition.interimResults = false;
                recognition.lang = 'en-US';

                recognition.onstart = function() {
                    recognizing = true;
                    document.getElementById('micBtn').innerText = '🔴 Recording Voice... (Click to Finish)';
                    document.getElementById('micBtn').style.backgroundColor = '#DC2626';
                    document.getElementById('status').innerText = 'Listening... Speak naturally.';
                };

                recognition.onresult = function(event) {
                    var transcript = '';
                    for (var i = event.resultIndex; i < event.results.length; ++i) {
                        transcript += event.results[i][0].transcript + ' ';
                    }
                    var textAreas = window.parent.document.querySelectorAll('textarea');
                    if (textAreas.length > 1) {
                        var target = textAreas[1];
                        var existing = target.value ? target.value.trim() + ' ' : '';
                        target.value = existing + transcript.trim();
                        target.dispatchEvent(new Event('input', { bubbles: true }));
                    }
                };

                recognition.onerror = function(event) {
                    document.getElementById('status').innerText = 'Mic notice: ' + event.error;
                    stopDictation();
                };

                recognition.onend = function() {
                    stopDictation();
                };
            } else {
                document.getElementById('status').innerText = 'Speech recognition not supported in browser.';
            }

            function toggleDictation() {
                if (recognizing) {
                    recognition.stop();
                    stopDictation();
                } else {
                    if (recognition) recognition.start();
                }
            }

            function stopDictation() {
                recognizing = false;
                document.getElementById('micBtn').innerText = '🎙️ Click to Speak Instructions';
                document.getElementById('micBtn').style.backgroundColor = '#2563EB';
                document.getElementById('status').innerText = 'Notes recorded. You can edit or add more text.';
            }
        </script>
        """,
        height=45,
    )

    special_instructions = st.text_area(
        "Voice or Typed Notes:",
        height=75,
        placeholder="Click mic above or type instructions here...",
        label_visibility="collapsed",
    )

    col_btn1, col_btn2 = st.columns([1, 1])
    with col_btn1:
        generate_btn = st.button("🚀 Generate Tailored Application Pack", type="primary", use_container_width=True)
    with col_btn2:
        gap_analyze_btn = st.button("🧠 Analyze JD for Knowledge Gaps", type="secondary", use_container_width=True)

    # ==============================================================================
    # SECTION 1: INTELLIGENT GAP QUESTIONING & ARCHIVE FEEDER
    # ==============================================================================
    if gap_analyze_btn:
        if not job_desc.strip():
            st.warning("Please paste a target Job Description first to analyze gaps.")
        else:
            with st.spinner("Analyzing JD against candidate master knowledge archive..."):
                full_archive_text = get_full_knowledge_context()
                gap_prompt = f"""
                You are a senior executive headhunter and career advisor for Madhusudhanan Janakarajan.
                Compare the target Job Description against the candidate's verified knowledge archive.

                CANDIDATE MASTER KNOWLEDGE ARCHIVE:
                {full_archive_text}

                TARGET JOB DESCRIPTION:
                {job_desc}

                TASK:
                1. Identify any specific technical platforms, niche distribution models, regional compliance, or specialized duties in the JD that are not clearly documented in the candidate's archive.
                2. Formulate 3 to 5 direct, precise questions asking the candidate if they have handled similar mandates (asking for company, scale, and concrete outcome).
                3. If the archive already completely answers the JD with high-conviction metrics, state that the profile has 95%+ direct proof points.

                Format output as a clean list of questions.
                """
                client = genai.Client(api_key=api_key)
                try:
                    gap_resp = client.models.generate_content(
                        model="gemini-3.6-flash",
                        contents=gap_prompt,
                        config=types.GenerateContentConfig(temperature=0.2)
                    )
                    st.session_state["gap_questions"] = gap_resp.text.strip()
                except Exception as e:
                    st.error(f"Gap Analysis Error: {str(e)}")

    if st.session_state.get("gap_questions"):
        st.markdown("---")
        st.markdown("#### ❓ Knowledge Archive Gap Analysis")
        st.info(st.session_state["gap_questions"])

        st.markdown("##### 📝 Feed Answers to Master Knowledge Base")
        st.caption("Paste your narrative answers below. They will be integrated and stored permanently to enrich this and all future applications.")
        new_answers = st.text_area("Your Narrative Answers / Verified Specifics:", height=110, placeholder="Example: At Conektr, we handled Red Bull under an 18-month exclusive terms agreement...")

        if st.button("💾 Save Answers to Permanent Master Archive"):
            if not new_answers.strip():
                st.warning("Please enter your answers before saving.")
            else:
                with st.spinner("Validating and adding new facts into archive..."):
                    # Conflict Resolution & Structuring pass
                    ingest_prompt = f"""
                    You are a knowledge base curator for Madhusudhanan Janakarajan.
                    Review the new facts provided by the candidate against their established background.
                    Ensure there are no contradictions (e.g. candidate manages FMCG/CPG and tech, not boutique luxury retail).
                    Structure these new points into concise, permanent facts.

                    NEW CANDIDATE INPUTS:
                    {new_answers}

                    Return a JSON array of objects, where each object has:
                    "q": "What requirement does this answer?",
                    "a": "Concise, verified fact with metrics and context"
                    """
                    client = genai.Client(api_key=api_key)
                    try:
                        ing_resp = client.models.generate_content(
                            model="gemini-3.6-flash",
                            contents=ingest_prompt,
                            config=types.GenerateContentConfig(response_mime_type="application/json", temperature=0.1)
                        )
                        parsed_new = json.loads(ing_resp.text.strip())
                        if save_custom_knowledge(parsed_new):
                            st.success(f"✅ Successfully added {len(parsed_new)} new verified facts to your permanent archive! You can now click 'Generate' to use them.")
                            del st.session_state["gap_questions"]
                            st.rerun()
                    except Exception as e:
                        st.error(f"Archive Update Error: {str(e)}")

# ==============================================================================
# MAIN GENERATION CONTROLLER (ROUTING + DEEP ARCHIVE SYNTHESIS)
# ==============================================================================
if generate_btn:
    if not job_desc or not job_desc.strip():
        st.warning("Please paste a target Job Description first.")
    elif not api_key:
        st.error("API Key is missing. Please configure GEMINI_API_KEY in Streamlit Secrets.")
    else:
        with col2:
            with st.spinner("⚡ High-Speed Synthesis: Applying contextual track routing and assembling documents..."):
                archive_context = get_full_knowledge_context()
                
                prompt = f"""
                You are an executive resume architect and career strategist for Madhusudhanan Janakarajan (23+ year FMCG, Digital Transformation & Enterprise Technology Executive).

                Analyze the provided target Job Description (JD) and special instructions. Cross-reference with the comprehensive candidate archive below to extract company name, role title, and generate fully customized documents.

                CANDIDATE MASTER KNOWLEDGE ARCHIVE:
                {archive_context}

                TARGET TRACK ROUTING RULES (DETECT AUTOMATICALLY FROM JD):
                - TRACK A (Commercial Leadership / RTM / Head of Sales / General Manager): Focus on $100M+ P&L, 6 GCC master distributors, 200+ India distributors, Net Sales Value (NSV), Gross Margin (GM), trade economics, drop-size thresholds, and multi-category scale (Britannia biscuits/dairy, Coca-Cola, Red Bull exclusivity, PepsiCo). Headers: "Traditional FMCG Operator" | "Digital FMCG Distribution" | "Distribution Transformation". Priority: commercial -> entrepreneurship -> transformation -> digital -> capability.
                - TRACK B (Commercial Excellence / Transformation / Capability / OD): Prioritize sales capability, Capability Champion networks, Train-the-Trainer (TTT), SPIN Selling, Business Measurement Index (BMI) distributor audits, 70:20:10 learning models, and macro Organization Design (OD). Headers: "Commercial Capability Leadership" | "Distributor Ops Excellence" | "Enterprise Tech Transformation". Priority: capability -> transformation -> commercial -> digital -> entrepreneurship.
                - TRACK C (Sales IT / Commercial IT / SFA-DMS Systems / Tech Solutions): Prioritize enterprise SaaS modernizations, Kenya on-ground P&G deployment, 1,000+ to 2,000+ sales user SFA rollouts, ERP-CRM-DMS integrations (SAP SD, Dynamics 365, Oracle), AI route optimization, Image Recognition (IR) audits, and analytics (Power BI, Alteryx, Tableau). Headers: "Core Commercial Operations" | "Digital Distribution Platforms" | "Enterprise Sales IT & SFA". Priority: transformation -> digital -> commercial -> capability -> entrepreneurship.
                - TRACK D (Digital Commerce / Marketplace / Omnichannel / Product Leadership): Prioritize Conektr founding/scale (8,000+ B2B stores), BOSS loyalty micro-hubs, beverage volume share (>50% Red Bull, Coca-Cola, waters, health drinks), full SDLC ownership ($27M platform valuation), Click & Collect / Scan & Go store integrations, 2-wheeler last-mile logistics, 60% WhatsApp ordering mix, and $15M M&A exit to Al Maya Group. Headers: "Core FMCG Operator" | "Digital Marketplace & B2B2C" | "RTM Modernization Advisory". Priority: digital -> entrepreneurship -> commercial -> transformation -> capability.

                STRICT EXECUTIVE WRITING RULES:
                - NEVER spell out numbers or metric notations into words. Always write "360°" (NEVER "thirty-six-degree" or "360-degree"), "$100M+" (NEVER "one hundred million"), "23+ years" (NEVER "twenty-three years"), "8,000+" (NEVER "eight thousand"), and "~40%" (NEVER "forty percent").
                - Synthesize casual or informal special instructions into polished, authoritative executive phrasing.

                JSON SCHEMA REQUIREMENTS:
                1. IDENTIFY TARGET COMPANY & ROLE:
                   - "target_company": Specific company name from JD.
                   - "target_role": Specific role title from JD.

                2. HEADER SUBTITLE DUAL VARIABLES:
                   - Format: "[header_focus_1] | FMCG | GTM & Omnichannel Leader | [header_focus_2]"
                   - "header_focus_1": Target leadership title matching JD. Max 36 chars.
                   - "header_focus_2": Specialized domain focus matching JD. Max 40 chars.

                3. EXECUTIVE SUMMARY (STRICTLY 155 TO 170 WORDS / EXACTLY 8 FULL JUSTIFIED LINES):
                   - Authoritative, high-impact, passionate executive summary of EXACTLY 155 to 170 words tailored directly to the detected track and target company.
                   - Must completely fill 8 full justified lines in Calibri 10pt (line spacing multiple 1.16).
                   - Deliver a compelling narrative covering 23+ years driving FMCG commercial strategy, 360° operational vantage, entrepreneurship, and measurable enterprise results.

                4. CAPABILITY ORDERING (PRIORITIZATION):
                   - Array of all 5 capability keys ordered by detected track priority.

                5. EXPERIENCE COLUMN HEADERS (3 COLUMNS):
                   - "exp_col_header_1": Title for Column 1 matching detected track.
                   - "exp_col_header_2": Title for Column 2 matching detected track.
                   - "exp_col_header_3": Title for Column 3 matching detected track.

                6. CONEKTR CATEGORY BULLET:
                   - Category aggregation bullet strictly tailored to products/domain of the target company.

                7. DYNAMIC EXPERIENCE INJECTIONS (STRICT 18 TO 24 WORDS EACH):
                   - "column_1_extra_bullet": 18-24 words under Britannia / Traditional FMCG regarding Route-to-Market, distributor governance, or commercial expansion aligned with JD, else "".
                   - "column_2_extra_bullet": 18-24 words under Conektr (Digital FMCG) aligned with JD, else "".
                   - "column_3_extra_bullet": 18-24 words under TransCPG/Ivy (Transformation) aligned with JD, else "".

                8. ATS MATCH SCORE (INTEGER 88-97):
                   - "ats_match_score": Integer reflecting alignment with provided JD.

                9. COVER LETTER & MATCH MATRIX:
                   - "subject_line": "Application for [Target Role] - [Target Company]"
                   - "cover_para_1": Authoritative opening explicitly referencing company name, role title, and candidate's 23+ year track record.
                   - "cover_para_2": Direct alignment with company's commercial/digital priorities based on JD & special instructions.
                   - "cover_bullets": 4 high-impact bullets formatted as "Bold Category: Detailed metric description":
                     1) Global Distributor Management & Commercial Governance: ...
                     2) Enterprise Digital Architecture & SFA Systems: ...
                     3) Measurable P&L & Operational ROI: ...
                     4) Cross-Functional Leadership & Partner Strategy: ...
                   - "cover_para_closing": Forward-looking closing paragraph.
                   - "matrix_items": Array of EXACTLY 6 rich, highly detailed competency rows mapping JD pillars to quantifiable candidate evidence.
                     * "requirement_title": Concise single statement of the requirement.
                     * "match_desc": Detailed, high-evidence paragraph with specific achievements, platforms, and metrics (without duplicate bold prefixes).

                INPUT JOB DESCRIPTION:
                {job_desc}

                INPUT SPECIAL INSTRUCTIONS / CONTEXT:
                {special_instructions}

                Return ONLY a valid JSON object matching this schema:
                {{
                  "target_company": "string",
                  "target_role": "string",
                  "header_focus_1": "string",
                  "header_focus_2": "string",
                  "executive_summary": "string",
                  "capability_order": ["string", "string", "string", "string", "string"],
                  "exp_col_header_1": "string",
                  "exp_col_header_2": "string",
                  "exp_col_header_3": "string",
                  "conektr_category_bullet": "string",
                  "column_1_extra_bullet": "string",
                  "column_2_extra_bullet": "string",
                  "column_3_extra_bullet": "string",
                  "ats_match_score": 94,
                  "cover_letter_data": {{
                    "target_company": "string",
                    "subject_line": "string",
                    "cover_para_1": "string",
                    "cover_para_2": "string",
                    "cover_bullets": ["string", "string", "string", "string"],
                    "cover_para_closing": "string",
                    "matrix_items": [
                      {{
                        "requirement_title": "string",
                        "match_desc": "string"
                      }}
                    ]
                  }}
                }}
                """

                client = genai.Client(api_key=api_key)
                tailored_data = None
                cover_data = None
                last_error = ""

                try:
                    response = client.models.generate_content(
                        model="gemini-3.6-flash",
                        contents=prompt,
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json", temperature=0.2
                        ),
                    )
                    raw_text = response.text.strip()
                    if raw_text.startswith("```"):
                        raw_text = re.sub(r"^```(?:json)?\s*", "", raw_text)
                        raw_text = re.sub(r"\s*```$", "", raw_text)

                    parsed_json = json.loads(raw_text)
                    parsed_json = sanitize_json_payload(parsed_json)

                    ordered_keys = parsed_json.get(
                        "capability_order",
                        ["commercial", "digital", "transformation", "capability", "entrepreneurship"],
                    )
                    full_capabilities = [BASE_CAPABILITIES[k] for k in ordered_keys if k in BASE_CAPABILITIES]
                    for k, cap_text in BASE_CAPABILITIES.items():
                        if cap_text not in full_capabilities:
                            full_capabilities.append(cap_text)

                    parsed_json["capabilities"] = full_capabilities
                    tailored_data = parsed_json
                    cover_data = parsed_json.get("cover_letter_data", {})
                except Exception as e:
                    last_error = str(e)

                if tailored_data:
                    st.session_state["tailored_data"] = tailored_data
                    st.session_state["cover_data"] = cover_data
                    st.session_state["job_desc"] = job_desc
                    rebuild_all_documents()
                    st.session_state["has_results"] = True
                else:
                    st.error(f"Generation Error: {last_error}. Please check your API key and network permissions.")

# ==============================================================================
# 6. PERSISTENT DISPLAY & SECTION 2: IN-PLACE REVISION ENGINE (3-DOCX SUITE)
# ==============================================================================
if st.session_state.get("has_results", False):
    with col2:
        tailored_data = st.session_state["tailored_data"]
        cover_data = st.session_state.get("cover_data", {})
        score = tailored_data.get("ats_match_score", 94)
        target_co = tailored_data.get("target_company", "Target Organization")
        target_rl = tailored_data.get("target_role", "Executive Position")

        st.subheader(f"2. Application Pack: {target_co}")
        st.caption(f"Role: **{target_rl}**")

        st.markdown(
            f"""
            <div style="
                display: flex;
                align-items: center;
                gap: 16px;
                padding: 12px 16px;
                background: #F0FDF4;
                border: 1px solid #BBF7D0;
                border-radius: 10px;
                margin-bottom: 15px;
            ">
                <div style="
                    width: 56px;
                    height: 56px;
                    border-radius: 50%;
                    background: conic-gradient(#16A34A {score * 3.6}deg, #E5E7EB 0deg);
                    display: flex;
                    align-items: center;
                    justify-content: center;
                ">
                    <div style="
                        width: 42px;
                        height: 42px;
                        border-radius: 50%;
                        background: white;
                        display: flex;
                        align-items: center;
                        justify-content: center;
                        font-weight: bold;
                        color: #16A34A;
                        font-size: 14px;
                    ">{score}%</div>
                </div>
                <div>
                    <div style="font-weight: 700; font-size: 14.5px; color: #166534;">Target JD Alignment Score: {score}/100</div>
                    <div style="font-size: 12px; color: #15803D;">High Match: Contextually tailored to {target_co}'s role specifications.</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # 1-Click Master ZIP (3 Files)
        st.download_button(
            label="📦 Download Application Bundle (.ZIP) — 3 Word Files",
            data=st.session_state["master_zip"],
            file_name=f"Madhusudhanan_Janakarajan_{target_co.replace(' ', '_')}_Application_Bundle.zip",
            mime="application/zip",
            type="primary",
            use_container_width=True,
        )

        st.markdown("---")
        st.write("📄 **Individual Application Files (.docx):**")

        st.download_button(
            label="📝 1. Complete Application Set (.docx) — Cover + Resume + Matrix",
            data=st.session_state["comb_docx"],
            file_name="1_Complete_Application_Set_Cover_Resume_Matrix.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        col_b1, col_b2 = st.columns(2)
        with col_b1:
            st.download_button(
                label="🟡 2. Highlighted Review Resume (.docx)",
                data=st.session_state["review_docx"],
                file_name="2_Madhusudhanan_Janakarajan_Resume_Highlighted_Review.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with col_b2:
            st.download_button(
                label="📄 3. Clean Master Resume (.docx)",
                data=st.session_state["clean_docx"],
                file_name="3_Madhusudhanan_Janakarajan_Resume_Clean.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        # ==============================================================================
        # SECTION 2: IN-PLACE LIVE CORRECTIONS & RE-RENDERING
        # ==============================================================================
        st.markdown("---")
        st.subheader("✍️ Instant Revisions / Feedback on Current Pack")
        st.caption("Request quick tweaks to this specific pack without re-pasting the JD. The AI will revise and update the downloads above in place.")

        correction_text = st.text_area(
            "Enter corrections or adjustments (e.g., 'Make match matrix points more detailed', 'Emphasize beverage distribution more in the summary'):",
            height=85,
            placeholder="Type your adjustments here...",
        )

        if st.button("🔄 Apply Revisions to Current Pack", type="secondary"):
            if not correction_text.strip():
                st.warning("Please type your feedback or correction first.")
            else:
                with st.spinner("⚡ Applying targeted corrections and rebuilding Word suite..."):
                    revise_prompt = f"""
                    You are refining an existing tailored executive application pack for Madhusudhanan Janakarajan.

                    CURRENT APPLICATION JSON DATA:
                    {json.dumps(st.session_state["tailored_data"])}

                    USER REVISION REQUEST / CORRECTION:
                    {correction_text}

                    STRICT REVISION RULES:
                    1. Apply the user's specific corrections directly to the relevant fields (e.g. executive_summary, matrix_items, cover_para, or bullet points).
                    2. Maintain all existing locked metrics and structures that were not asked to be changed.
                    3. NEVER write numbers as words. Ensure '360°', '$100M+', '23+ years', '8,000+', and '~40%' remain in numeric form.
                    4. Keep the executive_summary between 155 and 170 words (exactly 8 lines in 10pt Calibri).
                    5. Ensure matrix_items contain 6 detailed, metric-backed proof points.

                    Return ONLY the updated JSON with all fields intact.
                    """

                    client = genai.Client(api_key=api_key)
                    try:
                        rev_resp = client.models.generate_content(
                            model="gemini-3.6-flash",
                            contents=revise_prompt,
                            config=types.GenerateContentConfig(
                                response_mime_type="application/json", temperature=0.2
                            ),
                        )
                        rev_raw = rev_resp.text.strip()
                        if rev_raw.startswith("```"):
                            rev_raw = re.sub(r"^```(?:json)?\s*", "", rev_raw)
                            rev_raw = re.sub(r"\s*```$", "", rev_raw)

                        rev_json = json.loads(rev_raw)
                        rev_json = sanitize_json_payload(rev_json)

                        # Preserve full capabilities
                        ordered_keys = rev_json.get(
                            "capability_order",
                            ["commercial", "digital", "transformation", "capability", "entrepreneurship"],
                        )
                        full_caps = [BASE_CAPABILITIES[k] for k in ordered_keys if k in BASE_CAPABILITIES]
                        for k, cap_t in BASE_CAPABILITIES.items():
                            if cap_t not in full_caps:
                                full_caps.append(cap_t)
                        rev_json["capabilities"] = full_caps

                        st.session_state["tailored_data"] = rev_json
                        st.session_state["cover_data"] = rev_json.get("cover_letter_data", {})
                        rebuild_all_documents()
                        st.success("✅ Revisions successfully applied! Download the updated files above.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Failed to apply revision: {str(e)}")

        with st.expander("🔍 View AI Tailored Dynamic Variables"):
            st.write("**Identified Company:**", target_co)
            st.write("**Identified Role:**", target_rl)
            st.write("**Header Focus 1:**", tailored_data.get("header_focus_1"))
            st.write("**Header Focus 2:**", tailored_data.get("header_focus_2"))
            st.write("**Column 1 Header:**", tailored_data.get("exp_col_header_1"))
            st.write("**Column 2 Header:**", tailored_data.get("exp_col_header_2"))
            st.write("**Column 3 Header:**", tailored_data.get("exp_col_header_3"))
            st.write("**Executive Summary:**", tailored_data.get("executive_summary"))
            st.write("**Injected Bullet (Britannia/Traditional):**", tailored_data.get("column_1_extra_bullet"))
            st.write("**Injected Bullet (Conektr):**", tailored_data.get("column_2_extra_bullet"))
            st.write("**Injected Bullet (TransCPG/Ivy):**", tailored_data.get("column_3_extra_bullet"))
            st.write("**Match Matrix Rows Generated:**", len(cover_data.get("matrix_items", [])))
